import json
import argparse
import os
import zipfile
import shutil
from datetime import datetime, timezone

from docx import Document
from docx.oxml.ns import qn

# Global namespace map to avoid 'prefix a not found' errors
NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}


def is_list_paragraph(p):
    pPr = p._p.pPr
    if pPr is None:
        return False
    return pPr.numPr is not None


def get_list_level(p):
    pPr = p._p.pPr
    if pPr is None or pPr.numPr is None:
        return 0
    ilvl = pPr.numPr.ilvl
    if ilvl is None:
        return 0
    try:
        return int(ilvl.val)
    except Exception:
        return 0


def extract_paragraph_with_images(paragraph, image_map):
    """
    Extract text and inline images from a paragraph, returning a unified string.
    - Text stays as-is
    - Images become [[image:filename]]
    """
    runs_output = []

    for run in paragraph.runs:
        # Find all drawing elements (anchors/inline) in this run
        drawings = run._element.findall(".//w:drawing", namespaces=NSMAP)

        if drawings:
            for drawing in drawings:
                # Find the <a:blip> element that contains the r:embed ref
                blip = drawing.find(".//a:blip", namespaces=NSMAP)
                if blip is not None:
                    rid = blip.get(qn("r:embed"))
                    if rid in image_map:
                        image_filename = image_map[rid]
                        runs_output.append(f"[[image:{image_filename}]]")
            # Skip normal text in this run if we already handled images
            continue

        # Add normal text (if any)
        if run.text:
            runs_output.append(run.text)

    return "".join(runs_output).strip()


def extract_cell_text_preserve_lists_and_images(cell, image_map):
    """
    Build a string from a table cell, preserving:
    - bullet list indentation (\t)
    - inline image references ([[image:filename]])
    - normal text paragraphs
    """
    lines = []
    for p in cell.paragraphs:
        content = extract_paragraph_with_images(p, image_map)
        if not content:
            continue

        if is_list_paragraph(p):
            level = get_list_level(p)
            indent = "\t" * level
            lines.append(f"{indent}• {content}")
        else:
            lines.append(content)

    return "\n".join(lines)


def extract_cell_title_with_images(cell, image_map):
    """
    Extract title text from a cell, including image markers.
    No list formatting; we just join paragraph contents with spaces.
    """
    parts = []
    for p in cell.paragraphs:
        content = extract_paragraph_with_images(p, image_map)
        if content:
            parts.append(content)
    return " ".join(parts).strip()


def extract_images_from_docx(docx_path, output_folder, doc_basename):
    """
    Extract all images from a .docx file into output_folder.

    - Names follow: DOCFILENAME_Image00.ext
      e.g., 'Animal Care and Nursing_Image01.png'
    - Returns:
        image_map: dict mapping relationship IDs (rId) -> final filename.
        image_files: sorted list of all image filenames created for this doc.
    """
    doc = Document(docx_path)
    rels = doc.part.rels

    # Map relId -> target_ref (e.g. 'media/image1.png')
    relid_to_target = {}
    for rel in rels.values():
        if "image" in rel.reltype:
            relid_to_target[rel.rId] = rel.target_ref

    image_map = {}
    target_to_final = {}
    all_final_names = set()
    counter = 1  # per-document image numbering

    with zipfile.ZipFile(docx_path, "r") as z:
        for rid, target_ref in relid_to_target.items():
            # Normalize the path into the zip
            if target_ref.startswith("/"):
                target_ref = target_ref[1:]
            if not target_ref.startswith("word/"):
                zip_path = "word/" + target_ref
            else:
                zip_path = target_ref

            # If we've already extracted this underlying image file for this doc,
            # just reuse the final name.
            if zip_path in target_to_final:
                final_name = target_to_final[zip_path]
            else:
                # Use the original extension, but standardized basename
                _, orig_ext = os.path.splitext(zip_path)
                if not orig_ext:
                    orig_ext = ".bin"

                final_name = f"{doc_basename}_Image{counter:02d}{orig_ext}"
                counter += 1

                try:
                    with z.open(zip_path) as src, open(
                        os.path.join(output_folder, final_name), "wb"
                    ) as dst:
                        shutil.copyfileobj(src, dst)
                except KeyError:
                    # If the image is somehow missing, skip it gracefully
                    final_name = None

                target_to_final[zip_path] = final_name

            if final_name is not None:
                image_map[rid] = final_name
                all_final_names.add(final_name)

    return image_map, sorted(all_final_names)


def parse_docx_table_with_images(path, output_folder, doc_basename,
                                 table_index=0, has_header=True):
    """
    Parse a .docx file into JSON with images extracted.
    Returns (cards, image_files)
    """
    image_map, image_files = extract_images_from_docx(path, output_folder, doc_basename)
    doc = Document(path)

    if table_index >= len(doc.tables):
        raise IndexError(
            f"'{path}' has only {len(doc.tables)} tables, "
            f"but index={table_index} was requested."
        )

    table = doc.tables[table_index]
    cards = []
    start_row = 1 if has_header else 0

    for row in table.rows[start_row:]:
        cells = row.cells
        if len(cells) < 2:
            continue

        # Title may contain images now
        title = extract_cell_title_with_images(cells[0], image_map)
        # Details with bullets + images
        detail = extract_cell_text_preserve_lists_and_images(cells[1], image_map)

        if not title and not detail:
            continue

        cards.append({"title": title, "detail": detail})

    return cards, image_files


# ---------- Manifest helpers ----------

def load_manifest(manifest_path):
    """
    Load an existing manifest, returning:
      manifest_data: dict (raw JSON)
      index: dict[path] -> entry
    """
    if not os.path.exists(manifest_path):
        return {"generated_at": None, "files": []}, {}

    with open(manifest_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    index = {}
    for entry in data.get("files", []):
        path = entry.get("path")
        if path:
            index[path] = entry
    return data, index


def save_manifest(manifest_path, manifest_data, index):
    """
    Save manifest_data + index to manifest_path.
    """
    manifest_data["generated_at"] = datetime.now(timezone.utc).isoformat()
    manifest_data["files"] = sorted(
        index.values(),
        key=lambda e: e.get("path", "")
    )
    with open(manifest_path, "w", encoding="utf-8") as f:
        json.dump(manifest_data, f, ensure_ascii=False, indent=2)


def main():
    parser = argparse.ArgumentParser(
        description=(
            "Parse .docx files into JSON flashcard lists, "
            "preserving bullet lists, tab indentation, and images. "
            "Maintains a manifest.json for incremental updates."
        )
    )
    parser.add_argument(
        "-t", "--table-index",
        type=int,
        default=0,
        help="Index of the table to parse in each .docx (0-based, default: 0)"
    )
    parser.add_argument(
        "--no-header",
        action="store_true",
        help="Set this if your tables have no header row"
    )
    parser.add_argument(
        "--docx-dir",
        type=str,
        default=".",
        help="Directory to scan for .docx files (default: current working directory)"
    )
    parser.add_argument(
        "-o", "--output-dir",
        type=str,
        default=None,
        help=(
            "Directory where JSON, images, and manifest.json will be written. "
            "Default (if omitted): ../VTNECards/Data relative to docx-dir "
            "(for your legacy local workflow)."
        )
    )

    args = parser.parse_args()

    # Where to look for .docx files
    current_dir = os.path.abspath(args.docx_dir)

    # Where to write JSON / images / manifest
    if args.output_dir:
        output_dir = os.path.abspath(args.output_dir)
    else:
        # Preserve your old behavior when running locally without flags
        output_dir = os.path.abspath(os.path.join(current_dir, "..", "VTNECards", "Data"))

    os.makedirs(output_dir, exist_ok=True)

    manifest_path = os.path.join(output_dir, "manifest.json")
    manifest_data, manifest_index = load_manifest(manifest_path)

    docx_files = sorted(
        f for f in os.listdir(current_dir)
        if f.lower().endswith(".docx") and not f.startswith("~$")
    )

    if not docx_files:
        print(f"No .docx files found in {current_dir}.")
        return

    total_cards = 0

    for filename in docx_files:
        input_path = os.path.join(current_dir, filename)
        base, _ = os.path.splitext(filename)  # doc base name (with spaces)
        output_json_name = base + ".json"
        output_json_path = os.path.join(output_dir, output_json_name)

        print(f"\nProcessing {filename} ...")
        try:
            cards, image_files = parse_docx_table_with_images(
                path=input_path,
                output_folder=output_dir,
                doc_basename=base,
                table_index=args.table_index,
                has_header=not args.no_header,
            )
        except Exception as e:
            print(f"[ERROR] Failed to parse '{filename}': {e}")
            continue

        with open(output_json_path, "w", encoding="utf-8") as f:
            json.dump(cards, f, ensure_ascii=False, indent=2)

        print(f" → {len(cards)} cards → {output_json_path}")
        total_cards += len(cards)

        # --- Update manifest entries for this docx ---
        parsed_at = datetime.now(timezone.utc).isoformat()
        source_docx = filename

        # JSON entry
        json_entry = manifest_index.get(output_json_name, {"path": output_json_name})
        json_entry["type"] = "json"
        json_entry["source_docx"] = source_docx
        json_entry["parsed_at"] = parsed_at
        manifest_index[output_json_name] = json_entry

        # Image entries
        for img_name in image_files:
            img_entry = manifest_index.get(img_name, {"path": img_name})
            img_entry["type"] = "image"
            img_entry["source_docx"] = source_docx
            img_entry["parsed_at"] = parsed_at
            manifest_index[img_name] = img_entry

    save_manifest(manifest_path, manifest_data, manifest_index)

    print(
        f"\nDone. Exported a total of {total_cards} flashcards "
        f"from {len(docx_files)} file(s) into {output_dir}"
    )
    print(f"Manifest updated at: {manifest_path}")


if __name__ == "__main__":
    main()
